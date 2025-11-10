from werkzeug.utils import secure_filename
from flask import Flask, render_template, request, redirect, url_for, flash, session, send_file, g
from datetime import datetime
import os
import random
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io
import psycopg2
import psycopg2.extras
from werkzeug.security import generate_password_hash, check_password_hash
from hashlib import sha256  
import speech_recognition as sr
from pydub import AudioSegment
import tempfile
from fpdf import FPDF
import pandas as pd
from io import BytesIO
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from docx import Document       
from docx2pdf import convert    
import tempfile                          

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY')
app.config['DATABASE_URL'] = os.environ.get('DATABASE_URL')  
app.config['UPLOAD_FOLDER'] = '/tmp/uploads/'
# app.config['UPLOAD_FOLDER'] = 'static/uploads/'
app.config['ALLOWED_EXTENSIONS'] = {'wav', 'mp3', 'ogg'}
app.config['WHATSAPP_ADMIN'] = os.environ.get('WHATSAPP_ADMIN')
app.config['WHATSAPP_DEFAULT_MSG'] = 'Halo BMKG, saya ingin konfirmasi permohonan wawancara dengan token: '

# Buat folder unggahan
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

# Konfigurasi database PostgreSQL
def get_db():
    if 'db' not in g:
        # Dapatkan URL dari environment variable Vercel
        db_url = os.environ.get('DATABASE_URL')
        
        if not db_url:
            app.logger.error("DATABASE_URL environment variable not found")
            raise RuntimeError("Database configuration error")

        # Konversi URL format jika diperlukan
        if db_url.startswith('postgres://'):
            db_url = db_url.replace('postgres://', 'postgresql://', 1)

        # Tambahkan parameter SSL jika belum ada
        if 'sslmode=' not in db_url:
            db_url += '?sslmode=require'

        try:
            # Backup fungsi DNS asli
            import socket
            original_getaddrinfo = socket.getaddrinfo

            # Force IPv4 jika diperlukan
            def forced_ipv4_getaddrinfo(host, port, *args, **kwargs):
                try:
                    # Coba IPv4 dulu
                    return original_getaddrinfo(host, port, socket.AF_INET, *args, **kwargs)
                except socket.gaierror:
                    # Fallback ke IPv6 jika IPv4 gagal
                    return original_getaddrinfo(host, port, socket.AF_INET6, *args, **kwargs)

            socket.getaddrinfo = forced_ipv4_getaddrinfo

            # Buat koneksi dengan timeout dan konfigurasi khusus
            g.db = psycopg2.connect(
                db_url,
                connect_timeout=10,
                keepalives=1,
                keepalives_idle=30,
                keepalives_interval=10,
                keepalives_count=5
            )

            # Test koneksi langsung
            with g.db.cursor() as cur:
                cur.execute("SELECT 1")
                if cur.fetchone()[0] != 1:
                    raise RuntimeError("Database connection test failed")

            app.logger.info("Database connection established successfully")

        except psycopg2.OperationalError as e:
            app.logger.error(f"Operational Error: {str(e)}")
            # Restore DNS resolver asli
            socket.getaddrinfo = original_getaddrinfo
            raise RuntimeError("Could not connect to database. Please try again later.")
            
        except Exception as e:
            app.logger.error(f"Unexpected database error: {str(e)}")
            raise RuntimeError("Database service unavailable")

    return g.db

@app.teardown_appcontext
def close_connection(exception):
    db = g.pop('db', None)
    if db is not None:
        db.close()

def init_db():
    try:
        db = get_db()
        cursor = db.cursor()
        
        # Perbaikan sintaks untuk PostgreSQL
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id SERIAL PRIMARY KEY,
            username TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            role TEXT NOT NULL
        )
        ''')
        
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS interview_requests (
            id SERIAL PRIMARY KEY,
            token TEXT UNIQUE NOT NULL,
            interviewer_name TEXT NOT NULL,
            media_name TEXT NOT NULL,
            topic TEXT NOT NULL,
            method TEXT NOT NULL,
            datetime TEXT NOT NULL,
            meeting_link TEXT,
            status TEXT DEFAULT 'Pending',
            request_date TEXT NOT NULL,
            whatsapp_link TEXT
        )
        ''')
        
        cursor.execute('''
        CREATE TABLE IF NOT EXISTS audio_recordings (
            id SERIAL PRIMARY KEY,
            token TEXT NOT NULL,
            interviewee TEXT NOT NULL,
            interviewer TEXT NOT NULL,
            date TEXT NOT NULL,
            filename TEXT NOT NULL,
            transcript TEXT
        )
        ''')
        
        # Insert admin user dengan password yang di-hash dengan benar
        admin_password = os.environ.get('ADMIN_PASSWORD', 'password123')
        user_password = os.environ.get('USER_PASSWORD', 'user123')
        
        cursor.execute('''
            INSERT INTO users (username, password, role) 
            VALUES (%s, %s, %s)
            ON CONFLICT (username) DO NOTHING
        ''', ('admin', generate_password_hash(admin_password, method='pbkdf2:sha256', salt_length=16), 'admin'))
        
        cursor.execute('''
            INSERT INTO users (username, password, role) 
            VALUES (%s, %s, %s)
            ON CONFLICT (username) DO NOTHING
        ''', ('user1', generate_password_hash(user_password, method='pbkdf2:sha256', salt_length=16), 'user'))
        
        db.commit()
        print("Database initialized successfully")
    except Exception as e:
        print(f"Error initializing database: {str(e)}")

# =============================================================== #
# Data contoh
weather_data = {
    'today': {
        'date': datetime.now().strftime('%A, %d %B %Y'),
        'condition': 'Partly Cloudy',
        'temp': '28°C',
        'humidity': '65%',
        'wind': '10 km/h',
        'outlook': 'Hari ini diperkirakan cerah berawan dengan kemungkinan hujan ringan di sore hari. Warga diimbau untuk tetap waspada terhadap perubahan cuaca.'
    },
    'tomorrow': {
        'condition': 'Rainy',
        'temp': '26°C',
        'outlook': 'Besok diperkirakan hujan dengan intensitas sedang. Warga diharapkan membawa payung atau jas hujan saat beraktivitas di luar ruangan.'
    }
}

warnings = [
    {'title': 'Peringatan Banjir', 'image': 'flood-warning.jpg', 'description': 'Waspada potensi banjir di daerah rendah'},
    {'title': 'Cuaca Ekstrem', 'image': 'storm-warning.jpg', 'description': 'Prakiraan angin kencang dan hujan lebat'},
    {'title': 'Gelombang Tinggi', 'image': 'wave-warning.jpg', 'description': 'Peringatan gelombang tinggi di pesisir'}
]

press_releases = [
    {
        'id': 1,
        'title': 'Konferensi Pers Bulanan BMKG',
        'date': '2023-06-15',
        'summary': 'BMKG mengadakan konferensi pers bulanan untuk membahas perkembangan iklim terkini.',
        'content': 'Dalam konferensi pers ini, BMKG akan memaparkan analisis kondisi cuaca dan iklim selama sebulan terakhir serta prakiraan untuk bulan depan...'
    },
    {
        'id': 2,
        'title': 'Peluncuran Sistem Peringatan Dini Baru',
        'date': '2023-05-28',
        'summary': 'BMKG meluncurkan sistem peringatan dini terbaru dengan teknologi canggih.',
        'content': 'Sistem baru ini mampu mendeteksi potensi bencana meteorologi lebih cepat dan akurat...'
    }
]

# =============================================================== #
# Routes
@app.route('/')
def index():
    return render_template('index.html', 
                         press_releases=press_releases[:3],
                         weather=weather_data['today'],
                         warnings=warnings)

@app.route('/press-release/<int:release_id>')
def press_release_detail(release_id):
    release = next((r for r in press_releases if r['id'] == release_id), None)
    if not release:
        flash('Press release not found', 'danger')
        return redirect(url_for('index'))
    return render_template('press_release_detail.html', release=release)

def generate_token():
    return ''.join(random.choices('ABCDEFGHJKLMNPQRSTUVWXYZ23456789', k=8))


def send_email_notification(token, recipient_email, request_data):
    """Mengirim email notifikasi tentang permohonan wawancara"""
    try:
        # Konfigurasi email
        smtp_server = os.environ.get('SMTP_SERVER', 'smtp.gmail.com')
        smtp_port = int(os.environ.get('SMTP_PORT', 587))
        smtp_username = os.environ.get('SMTP_USERNAME')
        smtp_password = os.environ.get('SMTP_PASSWORD')
        sender_email = os.environ.get('SENDER_EMAIL', 'noreply@bmkg.example.com')
        
        if not all([smtp_username, smtp_password, sender_email]):
            app.logger.warning("Email configuration not complete, skipping email notification")
            return False
        
        # Buat pesan email
        subject = f"Konfirmasi Permohonan Wawancara BMKG - Token: {token}"
        
        # HTML email content
        html = f"""
        <html>
            <body>
                <h2>Permohonan Wawancara BMKG</h2>
                <p>Terima kasih telah mengajukan permohonan wawancara dengan BMKG. Berikut detail permohonan Anda:</p>
                
                <table border="0" cellpadding="5">
                    <tr><td><strong>Token</strong></td><td>{token}</td></tr>
                    <tr><td><strong>Nama Pewawancara</strong></td><td>{request_data['interviewer_name']}</td></tr>
                    <tr><td><strong>Media</strong></td><td>{request_data['media_name']}</td></tr>
                    <tr><td><strong>Topik</strong></td><td>{request_data['topic']}</td></tr>
                    <tr><td><strong>Metode</strong></td><td>{request_data['method']}</td></tr>
                    <tr><td><strong>Waktu yang Diminta</strong></td><td>{request_data['datetime']}</td></tr>
                </table>
                
                <p>Status permohonan Anda saat ini: <strong>Pending</strong>. Kami akan menghubungi Anda segera setelah permohonan diproses.</p>
                
                <p>Anda dapat memeriksa status permohonan dengan menggunakan token di atas.</p>
                
                <p>Hormat kami,<br>Tim BMKG</p>
            </body>
        </html>
        """
        
        # Setup email
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = recipient_email
        msg['Subject'] = subject
        msg.attach(MIMEText(html, 'html'))
        
        # Kirim email
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_username, smtp_password)
            server.send_message(msg)
        
        return True
    
    except Exception as e:
        app.logger.error(f"Email sending error: {str(e)}")
        return False

@app.route('/request-interview', methods=['GET', 'POST'])
def request_interview():
    if request.method == 'POST':
        token = generate_token()
        method = request.form.get('method')
        email = request.form.get('email', '')

        # Ambil semua data form
        interviewer_name = request.form.get('interviewer_name')
        media_name = request.form.get('media_name')
        topic = request.form.get('topic')
        datetime_req = request.form.get('datetime')
        meeting_link = request.form.get('meeting_link', '')

        # Buat template pesan WhatsApp
        message = f"""Permohonan Wawancara BMKG
        Halo Admin BMKG 👋,
        Berikut detail permohonan wawancara yang baru saja diajukan:

        🧾 Token: *{token}*
        👤 Nama Pewawancara: *{interviewer_name}*
        🏢 Media: *{media_name}*
        🗣️ Topik Wawancara: *{topic}*
        📅 Tanggal & Waktu: *{datetime_req}*
        💬 Metode: *{method.capitalize()}*
        🔗 Meeting Link (jika ada): {meeting_link if meeting_link else '-'}
        
        Mohon tindak lanjut konfirmasi jadwal.
        Terima kasih 🙏
        """

        # Encode pesan agar formatnya aman di URL WhatsApp
        from urllib.parse import quote
        encoded_message = quote(message)

        # Buat link WhatsApp final
        whatsapp_link = f"https://wa.me/{app.config['WHATSAPP_ADMIN']}?text={encoded_message}"

        # Simpan ke database
        request_data = {
            'token': token,
            'interviewer_name': interviewer_name,
            'media_name': media_name,
            'topic': topic,
            'method': method,
            'datetime': datetime_req,
            'meeting_link': meeting_link,
            'whatsapp_link': whatsapp_link,
            'request_date': datetime.now().strftime('%Y-%m-%d %H:%M'),
            'status': 'Pending'
        }

        try:
            db = get_db()
            cursor = db.cursor()
            cursor.execute('''
                INSERT INTO interview_requests 
                (token, interviewer_name, media_name, topic, method, datetime, meeting_link, whatsapp_link, request_date, status)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ''', (
                request_data['token'],
                request_data['interviewer_name'],
                request_data['media_name'],
                request_data['topic'],
                request_data['method'],
                request_data['datetime'],
                request_data['meeting_link'],
                request_data['whatsapp_link'],
                request_data['request_date'],
                request_data['status']
            ))
            db.commit()

            # Kirim email jika ada
            if email:
                send_email_notification(token, email, request_data)

            # Arahkan ke WhatsApp langsung (untuk semua metode)
            return redirect(whatsapp_link)

        except Exception as e:
            flash(f'Database error: {str(e)}', 'danger')
            app.logger.error(f'Database error: {str(e)}')

    return render_template('request_interview.html')


@app.route('/historical-data', methods=['GET', 'POST'])
def historical_data_view():
    if 'user' not in session:
        flash('Please login to view this page', 'danger')
        return redirect(url_for('login'))
    
    try:
        db = get_db()
        cursor = db.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        # Handle search/filter
        search_query = request.args.get('search', '')
        status_filter = request.args.get('status', 'all')
        date_from = request.args.get('date_from', '')
        date_to = request.args.get('date_to', '')
        
        base_query = '''
            SELECT ir.*, ar.filename as has_recording
            FROM interview_requests ir
            LEFT JOIN audio_recordings ar ON ir.token = ar.token
        '''
        
        conditions = []
        params = []
        
        if search_query:
            conditions.append('''
                (ir.interviewer_name ILIKE %s OR 
                ir.media_name ILIKE %s OR 
                ir.topic ILIKE %s OR 
                ir.token ILIKE %s)
            ''')
            params.extend([f'%{search_query}%'] * 4)
        
        if status_filter != 'all':
            conditions.append('ir.status = %s')
            params.append(status_filter)
        
        if date_from:
            conditions.append('ir.request_date >= %s')
            params.append(date_from)
        
        if date_to:
            conditions.append('ir.request_date <= %s')
            params.append(date_to)
        
        if conditions:
            base_query += ' WHERE ' + ' AND '.join(conditions)
        
        base_query += ' ORDER BY ir.request_date DESC'
        
        cursor.execute(base_query, params)
        historical_data = cursor.fetchall()
        
        return render_template(
            'historical_data.html',
            historical_data=historical_data,
            search_query=search_query,
            status_filter=status_filter,
            date_from=date_from,
            date_to=date_to
        )
    
    except Exception as e:
        flash(f'Database error: {str(e)}', 'danger')
        app.logger.error(f'Database error: {str(e)}')
        return redirect(url_for('index'))

@app.route('/export-data')
def export_data():
    if 'user' not in session or session.get('role') != 'admin':
        flash('Unauthorized access', 'danger')
        return redirect(url_for('login'))
    
    try:
        db = get_db()
        
        # Ambil data dari database
        query = '''
            SELECT 
                ir.token,
                ir.interviewer_name,
                ir.media_name,
                ir.topic,
                ir.method,
                ir.datetime as scheduled_time,
                ir.status,
                ir.request_date,
                ar.interviewee,
                ar.date as recording_date
            FROM interview_requests ir
            LEFT JOIN audio_recordings ar ON ir.token = ar.token
            ORDER BY ir.request_date DESC
        '''
        
        df = pd.read_sql(query, db)
        
        # Buat file Excel di memori
        output = BytesIO()
        writer = pd.ExcelWriter(output, engine='xlsxwriter')
        df.to_excel(writer, sheet_name='Data Wawancara', index=False)
        
        # Formatting
        workbook = writer.book
        worksheet = writer.sheets['Data Wawancara']
        
        # Set column widths
        worksheet.set_column('A:A', 15)  # Token
        worksheet.set_column('B:B', 20)  # Interviewer
        worksheet.set_column('C:C', 20)  # Media
        worksheet.set_column('D:D', 30)  # Topic
        worksheet.set_column('E:E', 15)  # Method
        worksheet.set_column('F:F', 20)  # Scheduled Time
        worksheet.set_column('G:G', 15)  # Status
        worksheet.set_column('H:H', 20)  # Request Date
        worksheet.set_column('I:I', 20)  # Interviewee
        worksheet.set_column('J:J', 20)  # Recording Date
        
        writer.close()
        output.seek(0)
        
        return send_file(
            output,
            as_attachment=True,
            download_name='data_wawancara_bmkg.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    
    except Exception as e:
        flash(f'Error exporting data: {str(e)}', 'danger')
        app.logger.error(f'Export error: {str(e)}')
        return redirect(url_for('historical_data_view'))

def transcribe_audio(audio_path):
    """Mengkonversi audio ke teks menggunakan Google Speech Recognition"""
    try:
        # Konversi ke format WAV jika perlu
        if audio_path.endswith('.mp3'):
            sound = AudioSegment.from_mp3(audio_path)
            wav_path = audio_path.replace('.mp3', '.wav')
            sound.export(wav_path, format="wav")
            audio_path = wav_path
        
        recognizer = sr.Recognizer()
        with sr.AudioFile(audio_path) as source:
            audio_data = recognizer.record(source)
            text = recognizer.recognize_google(audio_data, language="id-ID")
            return text
    except Exception as e:
        app.logger.error(f"Transcription error: {str(e)}")
        return None

# Modifikasi route recorder untuk menangani KEDUA alur kerja
@app.route('/recorder', methods=['GET', 'POST'])
def recorder():
    if 'user' not in session:
        flash('Please login to view this page', 'danger')
        return redirect(url_for('login'))
    
    try:
        db = get_db()
        cursor = db.cursor(cursor_factory=psycopg2.extras.DictCursor) # Gunakan DictCursor
        cursor.execute('SELECT role FROM users WHERE username = %s', (session['user'],))
        user = cursor.fetchone()
        
        if not user or user['role'] != 'admin':
            flash('Unauthorized access', 'danger')
            return redirect(url_for('index'))
    except Exception as e:
        flash(f'Database error: {str(e)}', 'danger')
        app.logger.error(f'Database error: {str(e)}')
        return redirect(url_for('index'))

    # --- LOGIKA POST (dari alur kerja App 2 Anda, biarkan saja) ---
    if request.method == 'POST':
        if 'audio_file' not in request.files:
            flash('No audio file uploaded', 'danger')
            return redirect(request.url)
        
        file = request.files['audio_file']
        if file.filename == '':
            flash('No selected file', 'danger')
            return redirect(request.url)
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            transcript = transcribe_audio(filepath) # Fungsi STT Anda
            
            try:
                cursor.execute('''
                    INSERT INTO audio_recordings 
                    (token, interviewee, interviewer, date, filename, transcript)
                    VALUES (%s, %s, %s, %s, %s, %s)
                ''', (
                    request.form.get('token'),
                    request.form.get('interviewee'),
                    request.form.get('interviewer'),
                    datetime.now().strftime('%Y-%m-%d %H:%M'),
                    filename,
                    transcript or request.form.get('transcript', '')
                ))
                db.commit()
                flash('Recording saved successfully!', 'success')
            except Exception as e:
                db.rollback()
                flash(f'Gagal menyimpan ke DB: {str(e)}', 'danger')

            return redirect(url_for('recorder'))
    
    # --- LOGIKA GET (MODIFIKASI DI SINI) ---
    # Template baru (App 1) butuh 'tokens'
    # Template lama (App 2) butuh 'recordings'
    # Kita kirim keduanya
    
    tokens = []
    recordings = []
    
    try:
        # 1. Ambil tokens untuk dropdown (INI YANG BARU)
        cursor.execute('SELECT token FROM interview_requests ORDER BY request_date DESC')
        tokens_data = cursor.fetchall()
        tokens = [{'token': row['token']} for row in tokens_data] # Sesuaikan format

        # 2. Ambil recordings (ini dari kode asli Anda)
        cursor.execute('SELECT * FROM audio_recordings ORDER BY date DESC')
        recordings = cursor.fetchall()

    except Exception as e:
        flash(f'Database error saat mengambil data: {str(e)}', 'danger')
        app.logger.error(f'Database error: {str(e)}')
    
    # Render template dengan KEDUA variabel
    return render_template('recorder.html', tokens=tokens, recordings=recordings)

@app.route('/generate-pdf/<int:recording_id>')
def generate_pdf(recording_id):
    if 'user' not in session:
        flash('Unauthorized access', 'danger')
        return redirect(url_for('login'))

    try:
        db = get_db()
        # Gunakan DictCursor untuk mengakses data berdasarkan nama kolom
        cursor = db.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        # PERBAIKAN: Gunakan SQL JOIN untuk mengambil data lengkap
        # dari kedua tabel (audio_recordings dan interview_requests)
        cursor.execute('''
            SELECT
                ar.token,
                ar.interviewee,
                ar.interviewer,
                ar.date as recording_date,
                ar.transcript,
                ir.media_name,
                ir.topic,
                ir.method,
                ir.datetime as schedule_time
            FROM audio_recordings ar
            LEFT JOIN interview_requests ir ON ar.token = ir.token
            WHERE ar.id = %s
        ''', (recording_id,))
        
        recording = cursor.fetchone()
        
        if not recording:
            flash('Recording not found', 'danger')
            return redirect(url_for('recorder'))

        # --- Implementasi Template DOCX (Seperti Kode Lama) ---

        # 1. Tentukan path template. Pastikan file ini ada di Vercel.
        #    Tempatkan 'template.docx' di folder 'static'
        base_dir = os.path.dirname(os.path.abspath(__file__))
        doc_path = os.path.join(base_dir, 'static', 'template.docx')
        
        if not os.path.exists(doc_path):
            flash('Template file not found.', 'danger')
            app.logger.error(f"Template file not found at {doc_path}")
            return redirect(url_for('recorder'))

        doc = Document(doc_path)

        # 2. Siapkan data pengganti (sesuaikan placeholder di .docx Anda)
        #    Kode lama Anda menggunakan {{waktu}}, {{jenis}}, dll.
        replacements = {
            '{{waktu}}': recording['schedule_time'] or '-',
            '{{jenis}}': recording['method'] or '-',
            '{{pewawancara}}': recording['interviewer'] or '-',
            '{{instansi}}': recording['media_name'] or '-',
            '{{narasumber}}': recording['interviewee'] or '-',
            '{{transkripsi}}': recording['transcript'] or '(Tidak ada transkrip)',
            '{{topik}}': recording['topic'] or '-'
        }

        # 3. Ganti placeholder di paragraf
        for p in doc.paragraphs:
            for key, val in replacements.items():
                if key in p.text:
                    # Ganti placeholder sambil mempertahankan style
                    inline = p.runs
                    for i in range(len(inline)):
                        if key in inline[i].text:
                            text = inline[i].text.replace(key, val)
                            inline[i].text = text

        # 4. Tentukan path sementara (Vercel HANYA bisa menulis ke /tmp)
        #    Kita akan menggunakan 'tempfile' untuk membuatnya
        with tempfile.TemporaryDirectory() as tmpdir:
            safe_token = recording['token'].replace(" ", "_")
            
            docx_filename = f"wawancara_{safe_token}.docx"
            pdf_filename = f"wawancara_{safe_token}.pdf"
            
            docx_path = os.path.join(tmpdir, docx_filename)
            pdf_path = os.path.join(tmpdir, pdf_filename)

            # 5. Simpan .docx yang sudah diisi ke /tmp
            doc.save(docx_path)
            
            # 6. Konversi .docx ke .pdf di /tmp
            try:
                convert(docx_path, pdf_path)
            except Exception as e:
                # Ini adalah bagian yang kemungkinan besar GAGAL di Vercel
                # karena tidak ada LibreOffice
                app.logger.error(f"docx2pdf conversion failed: {str(e)}")
                flash('Gagal mengonversi PDF. Dependensi server mungkin hilang.', 'danger')
                return redirect(url_for('recorder'))

            # 7. Baca file PDF dari /tmp ke buffer
            if os.path.exists(pdf_path):
                pdf_buffer = io.BytesIO()
                with open(pdf_path, 'rb') as f:
                    pdf_buffer.write(f.read())
                pdf_buffer.seek(0)
                
                # 8. Kirim file buffer ke pengguna
                return send_file(
                    pdf_buffer,
                    as_attachment=True,
                    download_name=pdf_filename,
                    mimetype='application/pdf'
                )
            else:
                raise FileNotFoundError("File PDF gagal dibuat setelah konversi.")

    except Exception as e:
        flash(f'Error generating PDF: {str(e)}', 'danger')
        app.logger.error(f'PDF generation error: {str(e)}')
        return redirect(url_for('recorder'))

@app.route('/upload_audio', methods=['POST'])
def upload_audio():
    """
    Menerima file audio 'rekaman.webm' dari MediaRecorder di frontend.
    Menyimpannya ke folder /tmp/uploads/ (sesuai app.config).
    """
    if 'user' not in session or session.get('role') != 'admin':
        return {'status': 'error', 'message': 'Unauthorized'}, 403

    if 'audio' not in request.files:
        return {'status': 'error', 'message': 'No audio file'}, 400
    
    audio = request.files['audio']
    ext = 'webm' # Sesuai JS kita

    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"rekaman_{timestamp}.{ext}"
    
    # Gunakan app.config['UPLOAD_FOLDER'] (yaitu /tmp/uploads/)
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    try:
        audio.save(path)
        # Di aplikasi produksi, Anda akan upload 'path' ini ke Supabase Storage
        app.logger.info(f"Audio disimpan sementara di: {path}")
        return {'status': 'ok', 'filename': filename}
    except Exception as e:
        app.logger.error(f"Gagal menyimpan audio: {str(e)}")
        return {'status': 'error', 'message': 'Server error'}, 500

# ---

@app.route('/get_topik/<token>')
def get_topik(token):
    """
    Mengambil 'topic' dari Supabase berdasarkan 'token' yang dipilih 
    di dropdown. Mengembalikan JSON.
    """
    if 'user' not in session:
        return {'topik': ''}, 403

    topik = ''
    try:
        db = get_db()
        cursor = db.cursor(cursor_factory=psycopg2.extras.DictCursor)
        cursor.execute("SELECT topic FROM interview_requests WHERE token = %s", (token,))
        data = cursor.fetchone()
        if data:
            topik = data['topic']
        return {'topik': topik}
    except Exception as e:
        app.logger.error(f"Gagal get topik: {str(e)}")
        return {'topik': ''}, 500

# ---

@app.route('/save_pdf', methods=['POST'])
def save_pdf():
    """
    Menerima data form (token, narasumber, transkrip) dari 'App 1'.
    Mengambil data lain dari Supabase, membuat DOCX, mengonversi ke PDF,
    dan menyimpan hasilnya ke tabel 'audio_recordings'.
    """
    if 'user' not in session or session.get('role') != 'admin':
        return {'status': 'error', 'message': 'Unauthorized'}, 403

    db = get_db()
    cursor = db.cursor(cursor_factory=psycopg2.extras.DictCursor)
    
    try:
        data = request.form
        token = data['token']
        narasumber = data['narasumber']
        teks = data['transkripsi']
        topik = data['topik'] 

        # 1. Ambil data pelengkap dari Supabase
        cursor.execute("SELECT * FROM interview_requests WHERE token = %s", (token,))
        req_data = cursor.fetchone()
        
        if not req_data:
            return {'status': 'error', 'message': 'Token tidak ditemukan'}, 404
        
        pewawancara = req_data['interviewer_name']
        waktu = req_data['datetime']
        instansi = req_data['media_name']
        jenis = req_data['method']
        
        # 2. Implementasi Template DOCX (Sama seperti /generate-pdf Anda)
        base_dir = os.path.dirname(os.path.abspath(__file__))
        doc_path = os.path.join(base_dir, 'static', 'template.docx')
        
        if not os.path.exists(doc_path):
            app.logger.error(f"Template file not found at {doc_path}")
            return {'status': 'error', 'message': 'Template file not found'}, 500

        doc = Document(doc_path)
        
        replacements = {
            '{{waktu}}': waktu or '-',
            '{{jenis}}': jenis or '-',
            '{{pewawancara}}': pewawancara or '-',
            '{{instansi}}': instansi or '-',
            '{{narasumber}}': narasumber or '-',
            '{{transkripsi}}': teks or '(Tidak ada transkrip)',
            '{{topik}}': topik or '-'
        }

        for p in doc.paragraphs:
            for key, val in replacements.items():
                if key in p.text:
                    inline = p.runs
                    for i in range(len(inline)):
                        if key in inline[i].text:
                            text = inline[i].text.replace(key, val)
                            inline[i].text = text
        
        # 3. Simpan & Konversi PDF di /tmp/uploads/
        pdf_folder = app.config['UPLOAD_FOLDER'] 
        
        safe_narasumber = narasumber.replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        pdf_filename = f"wawancara_{token}_{safe_narasumber}_{timestamp}.pdf"
        docx_filename = pdf_filename.replace(".pdf", ".docx")
        
        docx_path = os.path.join(pdf_folder, docx_filename)
        pdf_path = os.path.join(pdf_folder, pdf_filename)
        
        doc.save(docx_path)
        
        try:
            convert(docx_path, pdf_path)
        except Exception as e:
            app.logger.error(f"docx2pdf conversion failed: {str(e)}")
            return {'status': 'error', 'message': f'Gagal konversi PDF: {e}'}, 500
        
        # 4. Simpan hasil transkrip ke tabel 'audio_recordings'
        # Ini akan menimpa/membuat data yang terkait dengan token ini
        
        tgl_rekaman = datetime.now().strftime('%Y-%m-%d %H:%M')
        
        # Gunakan INSERT ... ON CONFLICT untuk UPSERT (update jika ada, insert jika tidak)
        cursor.execute("""
            INSERT INTO audio_recordings 
            (token, interviewee, interviewer, date, filename, transcript)
            VALUES (%s, %s, %s, %s, %s, %s)
            ON CONFLICT (token) DO UPDATE SET
                interviewee = EXCLUDED.interviewee,
                interviewer = EXCLUDED.interviewer,
                date = EXCLUDED.date,
                filename = EXCLUDED.filename,
                transcript = EXCLUDED.transcript
        """, (token, narasumber, pewawancara, tgl_rekaman, pdf_filename, teks))
        
        db.commit()

        # Hapus file docx sementara
        if os.path.exists(docx_path):
            os.remove(docx_path)
            
        app.logger.info(f"PDF disimpan sementara di: {pdf_path}")
        return {'status': 'ok', 'pdf_path': pdf_path}

    except Exception as e:
        db.rollback()
        app.logger.error(f'Gagal save PDF: {str(e)}')
        return {'status': 'error', 'message': str(e)}, 500

@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'user' in session:
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        try:
            db = get_db()
            cursor = db.cursor()
            cursor.execute('SELECT * FROM users WHERE username = %s', (username,))
            user = cursor.fetchone()
            
            # --- TAMBAHKAN LOGGING DI SINI ---
            app.logger.info(f"Mencoba login untuk user: {username}")
            app.logger.info(f"Data user dari DB: {user}")
            
            if user:
                # Cek sebelum hashing
                app.logger.info(f"Memeriksa hash: {user[2]}")
                if check_password_hash(user[2], password):
                    # Berdasarkan struktur tabel Anda:
                    # user[0] = id
                    # user[1] = username
                    # user[2] = password
                    # user[3] = role
                    session['user'] = user[1]  # Menyimpan username
                    session['role'] = user[3]  # MENYIMPAN ROLE PENGGUNA
                                    
                    flash('Login berhasil!', 'success')
                    return redirect(url_for('index'))
                else:
                    flash('Password salah', 'danger')
            else:
                flash('Username tidak ditemukan', 'danger')
                
        except Exception as e:
            # --- UBAH LOGGING ERROR DI SINI ---
            # Cetak error yang sebenarnya untuk debugging
            app.logger.error(f"LOGIN EXCEPTION: Terjadi error -> {e}", exc_info=True)
            flash('Terjadi kesalahan sistem', 'danger')
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('user', None)
    return redirect(url_for('index'))

# Inisialisasi database saat aplikasi dimulai
with app.app_context():
    init_db()

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)
